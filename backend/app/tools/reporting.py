"""@tool wrappers for the reporting domain.

Each wrapper is a thin LLM adapter: parse args, call services/domains/reporting,
shape JSON. Wire shapes preserve the legacy ``langchain_tools.py`` payloads so
the existing agent test suite continues to exercise this layer unchanged.
"""
from __future__ import annotations

import base64
import html
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Literal

from langchain_core.tools import tool
from pydantic import AliasChoices, BaseModel, Field

from app.schemas import PortfolioPositionSpec, PricingEnvironmentSnapshot
from app.services.deep_agent.capability_gate import capability_gated
from app.services.deep_agent.envelopes import ToolGroup
from app.services.domains import reporting as reporting_svc
from app.services.domains import risk as risk_svc

from ._shaping import shape_report_full, shape_report_row
from ._product_inputs import ToolPositionSnapshotSpec


# ----- args schemas -----------------------------------------------------------


class PortfolioSnapshotInput(BaseModel):
    positions: list[ToolPositionSnapshotSpec] = Field(default_factory=list)
    market: PricingEnvironmentSnapshot = Field(
        default_factory=PricingEnvironmentSnapshot
    )


class ReportBatchInput(BaseModel):
    title: str = "Desk Risk Report"
    report_type: str = "portfolio"
    portfolio: PortfolioSnapshotInput = Field(default_factory=PortfolioSnapshotInput)


class CreateReportInput(BaseModel):
    portfolio_id: int
    report_type: Literal["portfolio", "risk"] = "portfolio"
    title: str = "Agent Generated Desk Report"
    pricing_parameter_profile_id: int | None = Field(
        default=None,
        validation_alias=AliasChoices(
            "pricing_parameter_profile_id", "pricing_profile_id"
        ),
        description=(
            "Pricing parameter profile id to bind portfolio/risk report "
            "calculations to. Use the same id selected for the risk/pricing run."
        ),
    )


class ListReportsInput(BaseModel):
    portfolio_id: int | None = Field(
        default=None, description="Filter to one portfolio_id (from request_payload)."
    )
    report_type: Literal["portfolio", "risk", "rfq"] | None = Field(
        default=None, description="Filter by report type."
    )
    status: (
        Literal["queued", "running", "completed", "completed_with_errors", "failed"]
        | None
    ) = Field(default=None, description="Filter by job status.")
    limit: int = Field(default=20, ge=1, le=100, description="Max rows to return.")


class GetReportInput(BaseModel):
    report_id: int = Field(description="ReportJob id from list_reports.")


class WriteReportArtifactInput(BaseModel):
    title: str = Field(description="Report title used for the artifact title/name.")
    format: Literal["markdown", "docx", "html"] = Field(
        description="Artifact format to create."
    )
    body_markdown: str = Field(description="Canonical report body in Markdown.")
    body_html: str | None = Field(
        default=None,
        description="Optional pre-rendered HTML body. Used only for format='html'.",
    )
    filename_stem: str | None = Field(
        default=None,
        description="Optional filename stem before timestamp and extension.",
    )


# ----- tools ------------------------------------------------------------------


@capability_gated(group=ToolGroup.DOMAIN_READ)
@tool("run_report_batch", args_schema=ReportBatchInput)
def run_report_batch_tool(
    title: str, report_type: str, portfolio: PortfolioSnapshotInput
) -> dict[str, Any]:
    """Prepare a report batch payload and summary for audited report generation."""
    legacy_positions = [
        PortfolioPositionSpec.model_validate(position.to_legacy_payload())
        for position in portfolio.positions
    ]
    risk = risk_svc.calculate_risk(
        positions=legacy_positions,
        market=portfolio.market,
    )
    return reporting_svc.run_batch(
        title=title,
        report_type=report_type,
        risk_summary=risk.get("totals") or {},
    )


@capability_gated(group=ToolGroup.DOMAIN_READ)
@tool("list_reports", args_schema=ListReportsInput)
def list_reports_tool(
    portfolio_id: int | None = None,
    report_type: str | None = None,
    status: str | None = None,
    limit: int = 20,
) -> dict[str, Any]:
    """Return recent ReportJob rows, newest-first, with optional filters."""
    rows = reporting_svc.list_reports(
        portfolio_id=portfolio_id,
        report_type=report_type,  # type: ignore[arg-type]
        status=status,  # type: ignore[arg-type]
        limit=limit,
    )
    reports = [shape_report_row(job) for job in rows]
    return {"reports": reports, "total": len(reports)}


@capability_gated(group=ToolGroup.DOMAIN_READ)
@tool("get_report", args_schema=GetReportInput)
def get_report_tool(report_id: int) -> dict[str, Any]:
    """Return full ReportJob row for one id including artifact_paths and summary.

    Raises ValueError if the report_id is not found.
    """
    job = reporting_svc.get_report(report_id=report_id)
    if job is None:
        raise ValueError(f"Report job not found: report_id={report_id}")
    return shape_report_full(job)


@capability_gated(group=ToolGroup.DOMAIN_WRITE)
@tool("write_report_artifact", args_schema=WriteReportArtifactInput)
def write_report_artifact_tool(
    title: str,
    format: Literal["markdown", "docx", "html"],
    body_markdown: str,
    body_html: str | None = None,
    filename_stem: str | None = None,
) -> dict[str, Any]:
    """Create a downloadable thread-local report artifact.

    The chat layer materializes the returned ``artifacts`` entry under
    ``/api/artifacts/agent/thread-...``. This is intentionally separate from
    the legacy ``create_report`` job queue.
    """
    fmt = format.lower()
    stem = _safe_report_stem(filename_stem or title)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    extension = {"markdown": "md", "html": "html", "docx": "docx"}[fmt]
    virtual_path = f"/trading_desk/reports/{stem}_{timestamp}.{extension}"
    content_bytes, artifact = _report_artifact_payload(
        fmt=fmt,
        virtual_path=virtual_path,
        body_markdown=body_markdown,
        body_html=body_html,
    )
    return {
        "file_path": virtual_path,
        "format": fmt,
        "size_bytes": len(content_bytes),
        "artifacts": [artifact],
    }


@capability_gated(group=ToolGroup.DOMAIN_WRITE)
@tool("create_report", args_schema=CreateReportInput)
def create_report_tool(
    portfolio_id: int,
    report_type: str = "portfolio",
    title: str = "Agent Generated Desk Report",
    pricing_parameter_profile_id: int | None = None,
) -> dict[str, Any]:
    """Queue a persisted report job for a portfolio and audit the event."""
    return reporting_svc.create_report(
        portfolio_id=portfolio_id,
        report_type=report_type,  # type: ignore[arg-type]
        title=title,
        pricing_profile_id=pricing_parameter_profile_id,
    )


def _safe_report_stem(value: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip()).strip("._-")
    stem = re.sub(r"_+", "_", stem)
    return (stem or "report")[:80]


def _report_artifact_payload(
    *,
    fmt: str,
    virtual_path: str,
    body_markdown: str,
    body_html: str | None,
) -> tuple[bytes, dict[str, Any]]:
    if fmt == "markdown":
        content = body_markdown.encode("utf-8")
        return content, {
            "path": virtual_path,
            "size_bytes": len(content),
            "kind": "text",
            "content": body_markdown,
        }
    if fmt == "html":
        rendered = body_html if body_html is not None else _markdown_to_html(body_markdown)
        content = rendered.encode("utf-8")
        return content, {
            "path": virtual_path,
            "size_bytes": len(content),
            "kind": "text",
            "content": rendered,
        }
    if fmt == "docx":
        content = _markdown_to_docx_bytes(body_markdown)
        return content, {
            "path": virtual_path,
            "size_bytes": len(content),
            "kind": "binary",
            "content_b64": base64.b64encode(content).decode("ascii"),
        }
    raise ValueError(f"unsupported report artifact format: {fmt}")


def _markdown_to_html(markdown: str) -> str:
    lines: list[str] = [
        "<!doctype html>",
        '<html><head><meta charset="utf-8"></head><body>',
    ]
    in_ul = False
    in_ol = False
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            in_ul, in_ol = _close_html_lists(lines, in_ul, in_ol)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            in_ul, in_ol = _close_html_lists(lines, in_ul, in_ol)
            level = len(heading.group(1))
            lines.append(f"<h{level}>{html.escape(heading.group(2))}</h{level}>")
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            if in_ol:
                lines.append("</ol>")
                in_ol = False
            if not in_ul:
                lines.append("<ul>")
                in_ul = True
            lines.append(f"<li>{html.escape(bullet.group(1))}</li>")
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            if in_ul:
                lines.append("</ul>")
                in_ul = False
            if not in_ol:
                lines.append("<ol>")
                in_ol = True
            lines.append(f"<li>{html.escape(numbered.group(1))}</li>")
            continue
        in_ul, in_ol = _close_html_lists(lines, in_ul, in_ol)
        lines.append(f"<p>{html.escape(line)}</p>")
    _close_html_lists(lines, in_ul, in_ol)
    lines.append("</body></html>")
    return "\n".join(lines)


def _close_html_lists(lines: list[str], in_ul: bool, in_ol: bool) -> tuple[bool, bool]:
    if in_ul:
        lines.append("</ul>")
    if in_ol:
        lines.append("</ol>")
    return False, False


def _markdown_to_docx_bytes(markdown: str) -> bytes:
    from docx import Document

    doc = Document()
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.strip()
        if not line:
            idx += 1
            continue
        if _looks_like_table(lines, idx):
            idx = _add_markdown_table(doc, lines, idx)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2), level=level)
            idx += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            doc.add_paragraph(bullet.group(1), style="List Bullet")
            idx += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            doc.add_paragraph(numbered.group(1), style="List Number")
            idx += 1
            continue
        doc.add_paragraph(line)
        idx += 1
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _looks_like_table(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return "|" in lines[idx] and _is_table_separator(lines[idx + 1])


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_markdown_table(doc: Any, lines: list[str], idx: int) -> int:
    rows = [_table_cells(lines[idx])]
    idx += 2
    while idx < len(lines) and "|" in lines[idx].strip():
        rows.append(_table_cells(lines[idx]))
        idx += 1
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            table.cell(row_idx, col_idx).text = row[col_idx] if col_idx < len(row) else ""
    return idx


__all__ = [
    "run_report_batch_tool",
    "list_reports_tool",
    "get_report_tool",
    "write_report_artifact_tool",
    "create_report_tool",
]
